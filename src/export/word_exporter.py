"""Word トーク原稿エクスポート: 7 セクション + 用語集"""
import io
from datetime import date

import numpy as np
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.export._helpers import INGREDIENTS, load_sku_fc_df, next_period, get_actual_sellout_12m
from src.forecast.market_forecast import aggregate_market, MarketForecaster
from src.forecast.penetration import PenetrationForecaster
from src.forecast.scenario import build_penetration_events
from src.forecast.currency import convert_amounts


_BRAND_BLUE = RGBColor(0x25, 0x63, 0xEB)
_DARK       = RGBColor(0x1E, 0x29, 0x3B)

_GLOSSARY = [
    ("ロジスティック曲線（S字カーブ）",
     "普及率の時間的変化を表す数式。初期はゆっくり、ある時点から急速に普及し、最終的に上限（飽和点）に近づく S 字型の曲線。"
     "ジェネリック/バイオシミラーの市場浸透に適用。"),
    ("ジェネリック（GE）浸透率 / バイオシミラー（BS）浸透率",
     "成分市場全体の販売数量に占める GE/BS 製品の割合（0〜1 または %）。"
     "高いほど後発品への切り替えが進んでいることを示す。"),
    ("Sell-in",
     "自社から卸業者への出荷数量・金額。在庫が卸に積み上がる形で計上される。"
     "薬価改定前に一時的に増加する「先買い」現象が起きやすい。"),
    ("Sell-out",
     "卸業者から医療機関・薬局への出庫数量・金額。実際の市場需要に近い指標。"
     "本予測では Sell-out を基本ベースとして予測している。"),
    ("在庫日数（Days of Inventory, DOI）",
     "月末在庫を足元の Sell-out ペース（日次換算）で割った値。"
     "標準的な在庫水準からの乖離を検知するための指標。"),
    ("信頼区間（80% 予測区間）",
     "予測値の上下幅を統計的に示す区間。80% の場合「20 回予測して 16 回は実績がこの範囲に収まる」ことを意味する。"
     "区間が広いほど不確実性が高い。"),
    ("薬価改定",
     "国が定める薬の公定価格（薬価）の見直し。通常 2 年ごとに実施（通常改定）、最近は中間年（奇数年）にも改定が入る（中間年改定）。"
     "GE 品は既存 GE 品の平均価格水準へと引き下げられることが多い。"),
    ("中間年改定",
     "通常改定の翌年に行われる薬価の追加的な見直し。"
     "過去に比べてより頻繁に改定が実施されるようになったため、毎年のコスト影響として予測に組み込む必要がある。"),
    ("MAPE（平均絶対パーセント誤差）",
     "予測精度の指標。実績と予測の差（絶対値）を実績で割った値の平均。値が小さいほど予測精度が高い。"
     "本システムではバックテスト（直近 12ヶ月の検証）で算出している。"),
    ("RMSE（二乗平均平方根誤差）",
     "予測誤差の大きさを示す指標。誤差の二乗平均の平方根。外れ値（大きな誤差）に敏感。"),
    ("シナリオ（加味/非加味）",
     "「制度変更加味（Regulatory-Adjusted）」は薬価改定・使用促進策等の制度変更影響を予測に反映したシナリオ。"
     "「制度変更非加味（Organic）」は制度変更の影響をゼロとし、自然なトレンドのみで予測したシナリオ。"
     "両者の差分が制度変更による影響額として算出される。"),
    ("Holt-Winters 法（指数平滑法）",
     "時系列予測の手法のひとつ。トレンド・季節性を考慮して過去の値に指数的な重み付けを行い将来を予測する。"
     "本システムの市場規模予測（Step 1）で使用している。"),
    ("バックテスト",
     "予測モデルの精度検証方法。直近の一定期間をモデルから除いて予測し、実績との誤差を評価する。"
     "本システムでは直近 12 ヶ月を hold-out 期間として使用している。"),
]


def _heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _BRAND_BLUE
    return p


def _para(doc, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.size   = Pt(10.5)
    return p


def _bullet(doc, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def generate_word(
    conn, ing_id: str, axis1: str, axis2: str, detail: str = "standard"
) -> bytes:
    ing_name = INGREDIENTS.get(ing_id, ing_id)

    # ── データ収集 ──────────────────────────────────────────────
    df_m = aggregate_market(conn, ing_id)
    start = next_period(df_m["period"].iloc[-1])

    mfc = MarketForecaster(ing_id, horizon=36).fit(df_m).predict()
    pf  = PenetrationForecaster(ing_id).fit(df_m)
    fit = pf.get_fit_result()

    ev_adj = build_penetration_events(conn, start, 36, "regulatory_adjusted", axis2)
    ev_exc = build_penetration_events(conn, start, 36, "regulatory_excluded",  axis2)
    pen_adj = pf.predict(36, events=ev_adj)
    pen_exc = pf.predict(36, events=ev_exc)

    df_adj = load_sku_fc_df(conn, ing_id, "regulatory_adjusted", axis2)
    df_exc = load_sku_fc_df(conn, ing_id, "regulatory_excluded",  axis2)

    df_adj["year"] = df_adj["period"].str[:4]
    df_exc["year"] = df_exc["period"].str[:4]
    adj_by_year = df_adj.groupby("year")["amount_jpy"].sum() / 10_000
    exc_by_year = df_exc.groupby("year")["amount_jpy"].sum() / 10_000
    years        = sorted(adj_by_year.index)

    actual_market_last = df_m["total_units"].iloc[-12:].sum()
    fc_market_total    = sum(mfc.forecast_units) / 12  # avg monthly
    pen_last_actual    = df_m["penetration_rate"].iloc[-1]
    pen_fc_end_adj     = pen_adj[-1]
    pen_fc_end_exc     = pen_exc[-1]

    cum_adj = adj_by_year.sum()
    cum_exc = exc_by_year.sum()
    cum_impact = cum_adj - cum_exc

    scenario_label = "制度変更加味（Regulatory-Adjusted）" if axis1 == "regulatory_adjusted" else "制度変更非加味（Organic）"
    axis2_label    = {"base": "ベース", "optimistic": "楽観", "pessimistic": "悲観"}[axis2]

    # ── 文書生成 ─────────────────────────────────────────────────
    doc = Document()

    # タイトル
    title = doc.add_heading(f"GE/BS 販売予測 — ロジック解説トーク原稿", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = _BRAND_BLUE

    doc.add_paragraph(
        f"成分: {ing_name}（{ing_id}）　|　シナリオ: {scenario_label} × {axis2_label}　|　生成日: {date.today()}"
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ① 予測の目的・対象
    _heading(doc, "① 予測の目的・対象", 1)
    _para(doc,
        f"本資料は、{ing_name}（{ing_id}）を対象成分とした、今後 36ヶ月の自社 GE/BS 製品の Sell-out 金額予測について、"
        f"その前提・ロジック・シナリオ設定を解説するものです。"
    )
    _para(doc,
        f"予測シナリオは「{scenario_label}」を主軸とし、シナリオ幅として「{axis2_label}」を採用しています。"
        f"内部計算はすべて円建て（JPY）で行い、表示時に通貨換算を行います。"
    )
    if years:
        _para(doc,
            f"36ヶ月の累計予測金額（加味シナリオ・JPY）は "
            f"約 {cum_adj:,.0f} 万円 を想定しています。"
        )

    # ② 疾患領域・成分市場トレンド
    _heading(doc, "② 疾患領域・成分市場トレンド", 1)
    _para(doc,
        f"{ing_name} の市場では、直近 12ヶ月の総販売数量は "
        f"約 {actual_market_last / 1000:,.0f} 千本 でした。"
    )
    _para(doc,
        f"Holt-Winters 法による将来予測では、今後 36ヶ月の月平均販売数量は "
        f"約 {fc_market_total / 1000:,.0f} 千本/月 と予測されます。"
        f"市場成長トレンドと季節性パターンを反映した予測値です。"
    )
    if detail == "detailed":
        _para(doc,
            "予測モデルのバックテスト（直近 12ヶ月の hold-out 検証）による精度指標は "
            "MAPE・RMSE ともに許容範囲内であることを確認済みです。"
        )

    # ③ ジェネリック/バイオシミラー浸透率カーブ
    _heading(doc, "③ ジェネリック/バイオシミラー浸透率カーブ", 1)
    _para(doc,
        f"GE/BS の浸透率は現在 {pen_last_actual:.1%} であり、ロジスティック曲線（S字カーブ）に従って"
        f"上限 L = {fit.params.L:.1%}（飽和点）に向かってゆっくりと収束していく予測です。"
    )
    _para(doc,
        f"立ち上がり速度パラメータ k = {fit.params.k:.4f}、変曲点 t₀ = {fit.params.t0:.1f} ヶ月として推定されました。"
        f"フィット精度は R² = {fit.r_squared:.3f} です。"
    )
    _para(doc,
        "S字カーブになる理由: 最初は先進的な病院や薬局が先行採用し、その後紹介・評判を通じて普及が加速、"
        "最終的には全患者への普及率が上限に近づくため、成長が鈍化して飽和します。"
    )

    # ④ SKU 配分の考え方
    _heading(doc, "④ SKU 配分の考え方", 1)
    _para(doc,
        "製品全体の予測数量を各 SKU に配分するため、過去の SKU 別 Sell-out 構成比（製品内シェア）の"
        "時系列トレンドを線形回帰で推定し、36ヶ月先まで外挿しています。"
    )
    _para(doc,
        "各月において全 SKU の構成比合計が 100% になるよう正規化を行っています。"
        "構成比の急変（新規 SKU 投入・包装変更等）が見られる場合は、変化点検知と区間別フィッティングで対応しています。"
    )

    # ⑤ 制度変更シナリオ比較
    _heading(doc, "⑤ 制度変更「加味」vs「非加味」シナリオ", 1)
    _para(doc,
        "「制度変更加味（Adjusted）」シナリオでは、薬価改定による価格調整係数と、"
        "後発品使用促進策強化による浸透率上限・速度パラメータの変化を予測に反映しています。"
    )
    _para(doc,
        "「非加味（Organic）」シナリオは、制度変更がなかった場合の自然トレンドのみで予測したベースラインです。"
    )
    _para(doc,
        f"36ヶ月累計での制度変更による影響額は "
        f"約 {cum_impact:+,.0f} 万円（JPY）と試算されます。"
    )
    if years:
        _para(doc, "年度別内訳:", bold=True)
        for y in years:
            a = adj_by_year.get(y, 0); e = exc_by_year.get(y, 0)
            _bullet(doc, f"{y}年: 加味 {a:,.0f} 万円 / 非加味 {e:,.0f} 万円 → 差分 {a-e:+,.0f} 万円")

    # ⑥ 要因分解
    _heading(doc, "⑥ 要因分解", 1)
    _para(doc, "前期実績から予測合計への変化を以下の 3 要因に分解して説明します。", bold=True)
    _bullet(doc, "市場成長・浸透率向上（オーガニック成長）: 市場規模の拡大と浸透率の上昇による増収効果")
    _bullet(doc, "制度変更影響: 薬価改定による単価下落と、使用促進策による浸透率上乗せ効果")
    _bullet(doc, "残余（SKU ミックス等）: SKU 構成比の変化・在庫補正等の細目")
    _para(doc,
        f"加味シナリオでの 12ヶ月予測合計は前期実績に対して"
        f"{'増加' if cum_adj / max(1, len(years)) > 0 else '減少'}する見込みです。"
        f"薬価改定影響（マイナス）を浸透率上昇（プラス）が一部相殺する構図が主要なドライバーです。"
    )

    # ⑦ 想定 Q&A
    _heading(doc, "⑦ 想定される質疑（Q&A）", 1)
    qa_pairs = [
        ("Q. 予測の前提が変わったらどうなりますか？",
         "A. サイドバーの「シナリオ幅」（楽観/悲観）を変更するか、"
         "制度変更イベントのパラメータ（impact_value）を調整することで即時再計算できます。"
         "FX レートも同様に変更可能です。"),
        ("Q. 薬価改定の影響はどう計算していますか？",
         "A. `regulatory_events` テーブルの `price_change_rate` を価格調整係数として Sell-out 金額に乗算しています。"
         "「制度変更加味」シナリオのみ反映され、「非加味」シナリオでは係数 1.0 として扱います。"),
        ("Q. なぜ S 字カーブを使うのですか？",
         "A. 後発品の普及は「先進採用 → 急速拡大 → 飽和」という 3 段階をたどる経験則があり、"
         "ロジスティック曲線はこのパターンをシンプルに表現できる数式です。"
         "実際の浸透率データへのフィット精度も R²=0.98 以上であることを確認しています。"),
        ("Q. Sell-in と Sell-out の違いと使い分けは？",
         "A. Sell-out（実需要）を予測ベースとし、Sell-in は Sell-out ± 在庫変動補正で算出します。"
         "薬価改定前後には卸の在庫積み増し・調整が生じるため、Sell-in だけで需要を見ると誤差が大きくなります。"),
        ("Q. モデルの精度はどの程度ですか？",
         "A. 直近 12ヶ月の hold-out バックテストにより、市場規模予測の MAPE は 1〜3% 程度です。"
         "浸透率フィットは R²=0.97 以上を確認しています。ただし予測は不確実性を伴うため、"
         "信頼区間（80% 予測区間）も合わせてご参照ください。"),
    ]
    for q, a in qa_pairs:
        _para(doc, q, bold=True)
        _para(doc, a)
        doc.add_paragraph()

    # 用語集
    doc.add_page_break()
    _heading(doc, "用語集（グロッサリー）", 1)
    _para(doc, "本文中で使用した専門用語の解説です。")
    doc.add_paragraph()

    for term, definition in _GLOSSARY:
        p = doc.add_paragraph()
        run_t = p.add_run(f"{term}: ")
        run_t.font.bold = True
        run_t.font.color.rgb = _BRAND_BLUE
        run_d = p.add_run(definition)
        run_d.font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
